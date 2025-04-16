import streamlit as st
import pdfplumber
from docx import Document
import re
from datetime import datetime
from io import BytesIO
import pythoncom
import win32com.client as win32

# Configuration Streamlit
st.set_page_config(page_title="CV SCC Generator", layout="wide")
st.title("📄 CV SCC Generator 2024")

# Fonctions d'extraction
def extract_text_from_pdf(uploaded_file):
    with pdfplumber.open(uploaded_file) as pdf:
        return "\n".join([page.extract_text() for page in pdf.pages if page.extract_text()])

def extract_text_from_docx(uploaded_file):
    doc = Document(uploaded_file)
    return "\n".join([para.text for para in doc.paragraphs])

# Fonctions de traitement
def parse_personal_info(text):
    name = re.search(r"([A-Z][a-z]+ [A-Z][a-z]+)", text)
    return {
        'initials': '. '.join([n[0] for n in name.group().split()]) + '.' if name else "N.P.",
        'job_title': re.search(r"(?i)(technicien|administrateur|ingénieur).*", text).group() if re.search(r"(?i)(technicien|administrateur|ingénieur).*", text) else "Poste non spécifié"
    }

def parse_experience(text):
    experiences = re.findall(r"(?i)(\d{2}/\d{4}) - (\d{2}/\d{4}|présent)\s+(.*?)\s+(.*?)(?=\d{2}/\d{4}|$)", text, re.DOTALL)
    return [{
        'dates': f"{exp[0]} - {exp[1]}",
        'employer': exp[2].strip(),
        'position': exp[3].strip(),
        'missions': re.findall(r"• (.*?)(?=\n• |\n\w|$)", exp[-1])
    } for exp in experiences]

# Fonction de génération du DOCX
def generate_scc_docx(data):
    doc = Document('model_scc.docx')
    
    # Remplissage des champs
    for paragraph in doc.paragraphs:
        if "Prénom NOM" in paragraph.text:
            paragraph.text = paragraph.text.replace("Prénom NOM", data['initials'])
        
        if "Technicien Support de Proximité" in paragraph.text:
            paragraph.text = paragraph.text.replace("Technicien Support de Proximité", data['job_title'])
    
    return doc

# Interface Streamlit
uploaded_file = st.file_uploader("Déposez votre CV (PDF ou Word)", type=["pdf", "docx"])

if uploaded_file:
    if uploaded_file.type == "application/pdf":
        text = extract_text_from_pdf(uploaded_file)
    else:
        text = extract_text_from_docx(uploaded_file)
    
    data = {
        **parse_personal_info(text),
        'experiences': parse_experience(text),
        # Ajouter d'autres fonctions de parsing ici
    }
    
    docx = generate_scc_docx(data)
    
    # Sauvegarde et téléchargement
    output = BytesIO()
    docx.save(output)
    output.seek(0)
    
    st.success("CV généré avec succès !")
    st.download_button(
        label="📥 Télécharger le CV formaté",
        data=output,
        file_name="cv_scc_formaté.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
