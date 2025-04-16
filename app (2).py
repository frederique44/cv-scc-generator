
import streamlit as st
from docx import Document
from io import BytesIO
import pdfplumber

st.set_page_config(page_title="Générateur de CV SCC", layout="centered")

st.title("🧾 Générateur de CV format SCC")
st.write("Téléversez un fichier Word ou PDF avec les données du CV. Vous obtiendrez un CV formaté selon le modèle SCC.")

uploaded_file = st.file_uploader("📤 Téléverser un fichier Word (.docx) ou PDF (.pdf)", type=["docx", "pdf"])

def extract_data_from_text(text):
    data = {"certifications": [], "formations": [], "langues": [], "competences": [], "experiences": []}
    current = ""
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.lower().startswith("certifications"):
            current = "certifications"
        elif line.lower().startswith("formations"):
            current = "formations"
        elif line.lower().startswith("langues"):
            current = "langues"
        elif "compétence" in line.lower():
            current = "competences"
        elif "expérience" in line.lower():
            current = "experiences"
        elif current:
            data[current].append(line)
    return data

if uploaded_file:
    if uploaded_file.name.endswith(".docx"):
        doc_in = Document(uploaded_file)
        text = "\n".join([p.text for p in doc_in.paragraphs])
    elif uploaded_file.name.endswith(".pdf"):
        with pdfplumber.open(uploaded_file) as pdf:
            text = "\n".join([page.extract_text() for page in pdf if page.extract_text()])
    else:
        st.error("Format non pris en charge.")
        st.stop()

    data = extract_data_from_text(text)

    template = Document("Modèle_CV_SCC_2024.docx")

    for p in template.paragraphs:
        if "Prénom NOM" in p.text:
            p.text = "E. D."
        if "Technicien Support de Proximité" in p.text:
            p.text = "Technicien Systèmes et Réseaux"
        if "XX ans" in p.text:
            p.text = "30 ans d'expérience"

    template.add_page_break()
    template.add_paragraph("Données extraites :", style="Heading 2")
    for section, items in data.items():
        template.add_paragraph(section.capitalize(), style="Heading 3")
        for i in items:
            template.add_paragraph(f"• {i}", style="List Bullet")

    buffer = BytesIO()
    template.save(buffer)
    buffer.seek(0)

    st.success("✅ CV généré avec succès !")
    st.download_button("📥 Télécharger le CV SCC", buffer, file_name="CV_SCC_Généré.docx")
